import pytest
from docx import Document

from ingest.docx_extract_n_chunk import _heading_level, extract_chunks_from_doc


@pytest.mark.parametrize(
    "style, expected",
    [
        ("Heading 1", 1),
        ("Heading 3", 3),
        ("Normal", None),
        ("Heading", None),
        (None, None),
    ],
)
def test_heading_level(style, expected):
    assert _heading_level(style) == expected


def _build_doc():
    doc = Document()
    doc.add_paragraph("Scope", style="Heading 2")
    doc.add_paragraph("applies to media services")
    doc.add_paragraph("Details", style="Heading 3")
    doc.add_paragraph("more detail")
    doc.add_paragraph("Obligations", style="Heading 2")
    doc.add_paragraph("must exercise diligence")
    return doc


def test_extract_splits_on_headings():
    chunks = extract_chunks_from_doc(_build_doc(), title="Doc")
    assert len(chunks) == 3
    assert all(c["doc_title"] == "Doc" for c in chunks)


def test_extract_tracks_heading_hierarchy():
    chunks = extract_chunks_from_doc(_build_doc(), title="Doc")

    # Chunk under the H3 keeps its parent H2.
    assert chunks[1]["heading_2"] == "Scope"
    assert chunks[1]["heading_3"] == "Details"

    # A new H2 must drop the stale H3 from the deeper level.
    assert chunks[2]["heading_2"] == "Obligations"
    assert chunks[2]["heading_3"] == "na"


def test_blank_paragraphs_ignored():
    doc = Document()
    doc.add_paragraph("Head", style="Heading 2")
    doc.add_paragraph("   ")
    doc.add_paragraph("body")
    chunks = extract_chunks_from_doc(doc, title="Doc")
    assert len(chunks) == 1
    assert "body" in chunks[0]["chunk"]
